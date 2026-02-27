import os
import json
import pickle
import requests
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from tqdm import tqdm
import hashlib

# Файл уншигч сангууд
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️ PyPDF2 суугаагүй: pip install PyPDF2")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx суугаагүй: pip install python-docx")

try:
    import pandas as pd
    CSV_AVAILABLE = True
except ImportError:
    CSV_AVAILABLE = False
    print("⚠️ pandas суугаагүй: pip install pandas")

try:
    from pptx import Presentation
    PPTX_AVAILABLE = True
except ImportError:
    PPTX_AVAILABLE = False
    print("⚠️ python-pptx суугаагүй: pip install python-pptx")

load_dotenv()

# === CONFIG ===
EMBEDDING_MODEL = "sentence-transformers/all-distilroberta-v1"
INDEX_FOLDER = "faiss_disk_index"
METADATA_FILE = "disk_metadata.pkl"
SUPPORTED_EXTENSIONS = ['.txt', '.pdf', '.docx', '. doc', '.json', '.jsonl', '.csv', '.md', '.pptx', '.ppt']

# === OLLAMA CONFIG ===
OLLAMA_BASE_URL = "http://localhost:11434"
OLLAMA_MODEL = "qwen2"  # Таны татсан model


class OllamaLLM:
    """Ollama LLM wrapper"""
    
    def __init__(self, model: str = OLLAMA_MODEL, base_url: str = OLLAMA_BASE_URL):
        self.model = model
        self. base_url = base_url
        self.api_url = f"{base_url}/api/generate"
        
    def is_available(self) -> bool:
        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=10)
            return response.status_code == 200
        except:
            return False
    
    def list_models(self) -> list:
        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=10)
            if response.status_code == 200:
                data = response.json()
                return [m["name"] for m in data.get("models", [])]
            return []
        except:
            return []
    
    def generate(self, prompt: str, temperature: float = 0.3, max_tokens: int = 1024) -> str:
        try:
            payload = {
                "model": self.model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": temperature,
                    "num_predict": max_tokens,
                }
            }
            
            # 10 минут timeout
            response = requests.post(self.api_url, json=payload, timeout=600)
            
            if response.status_code == 200:
                data = response.json()
                return data.get("response", ""). strip()
            else:
                return f"Ollama error: {response.status_code}"
                
        except requests.exceptions. ReadTimeout:
            return "❌ Хариулт удаж байна (10 минутаас илүү).  Жижиг model ашиглана уу."
        except requests.exceptions.ConnectionError:
            return "❌ Ollama холбогдож чадсангүй."
        except Exception as e:
            return f"❌ Алдаа: {e}"
    
    def generate_stream(self, prompt: str, temperature: float = 0.3):
        try:
            payload = {
                "model": self.model,
                "prompt": prompt,
                "stream": True,
                "options": {
                    "temperature": temperature,
                }
            }
            
            # 10 минут timeout
            response = requests.post(self.api_url, json=payload, stream=True, timeout=600)
            
            for line in response. iter_lines():
                if line:
                    data = json. loads(line)
                    if "response" in data:
                        yield data["response"]
                    if data.get("done", False):
                        break
                        
        except requests.exceptions. ReadTimeout:
            yield "❌ Хариулт удаж байна."
        except Exception as e:
            yield f"❌ Алдаа: {e}"


class AdvancedDiskSearch:
    def __init__(self, search_paths=None):
        self.search_paths = search_paths or ["D:/", "C:/Users"]
        self.embedding = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
        self. db = None
        self.metadata = {}
        self.all_docs = []
        
        # Initialize Ollama instead of flan-t5
        print("🔄 Ollama холбогдож байна...")
        self.llm = OllamaLLM(model=OLLAMA_MODEL)
        
        if self.llm.is_available():
            models = self.llm.list_models()
            print(f"✅ Ollama бэлэн!  Моделууд: {', '.join(models[:5])}")
            
            # Check if selected model exists
            model_names = [m. split(":")[0] for m in models]
            if OLLAMA_MODEL not in model_names and OLLAMA_MODEL. split(":")[0] not in model_names:
                print(f"⚠️ '{OLLAMA_MODEL}' model олдсонгүй.")
                if models:
                    self.llm.model = models[0]
                    print(f"   '{self.llm.model}' model ашиглаж байна")
        else:
            print("❌ Ollama ажиллахгүй байна!")
            print("   1. Ollama суулгах: https://ollama.ai")
            print("   2.  Ажиллуулах: ollama serve")
            print("   3. Model татах: ollama pull qwen2")

    def get_file_hash(self, filepath):
        try:
            with open(filepath, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except:
            return None

    def should_skip_directory(self, dirpath):
        skip_dirs = {
            'node_modules', '__pycache__', '.git', '.venv', 'venv',
            'AppData', 'Windows', 'Program Files', 'System32',
            '$RECYCLE.BIN', 'Recovery', 'ProgramData',
            'Microsoft VS Code', 'Visual Studio', 'extensions',
            'resources', 'locales', 'vendor', 'build', 'dist'
        }
        return any(skip in dirpath for skip in skip_dirs)

    def read_txt_file(self, filepath):
        encodings = ['utf-8', 'utf-16', 'cp1252', 'latin-1']
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except:
                continue
        return None

    def read_pdf_file(self, filepath):
        if not PDF_AVAILABLE:
            return None
        try:
            text = ""
            with open(filepath, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text. strip()
        except Exception as e:
            print(f"⚠️ PDF уншихад алдаа {filepath}: {e}")
            return None

    def read_docx_file(self, filepath):
        if not DOCX_AVAILABLE:
            return None
        try:
            doc = docx. Document(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text. strip()
        except Exception as e:
            print(f"⚠️ DOCX уншихад алдаа {filepath}: {e}")
            return None

    def read_csv_file(self, filepath):
        if not CSV_AVAILABLE:
            return None
        try:
            df = pd.read_csv(filepath, encoding='utf-8', nrows=1000)
            return df.to_string()
        except Exception as e:
            try:
                df = pd.read_csv(filepath, encoding='latin-1', nrows=1000)
                return df.to_string()
            except:
                print(f"⚠️ CSV уншихад алдаа {filepath}: {e}")
                return None

    def read_json_file(self, filepath):
        try:
            if filepath.endswith('.jsonl'):
                texts = []
                with open(filepath, 'r', encoding='utf-8') as f:
                    for line in f:
                        data = json.loads(line)
                        texts.append(json.dumps(data, indent=2, ensure_ascii=False))
                return "\n---\n".join(texts[:100])
            else:
                with open(filepath, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"⚠️ JSON уншихад алдаа {filepath}: {e}")
            return None

    def read_pptx_file(self, filepath):
        if not PPTX_AVAILABLE:
            return None
        try:
            prs = Presentation(filepath)
            text = []
            for slide in prs. slides:
                for shape in slide.shapes:
                    if hasattr(shape, "text"):
                        text.append(shape.text)
            return "\n". join(text)
        except Exception as e:
            print(f"⚠️ PPTX уншихад алдаа {filepath}: {e}")
            return None

    def read_file(self, filepath):
        ext = Path(filepath).suffix.lower()
        if ext in ['.txt', '.md', '.log']:
            return self.read_txt_file(filepath)
        elif ext == '.pdf':
            return self. read_pdf_file(filepath)
        elif ext in ['.docx', '.doc']:
            return self.read_docx_file(filepath)
        elif ext == '.csv':
            return self.read_csv_file(filepath)
        elif ext in ['.json', '.jsonl']:
            return self.read_json_file(filepath)
        elif ext in ['.pptx', '.ppt']:
            return self.read_pptx_file(filepath)
        else:
            return None

    def scan_disk(self, max_files=1000, max_size_mb=10):
        print(f"🔍 Диск хайж эхэлж байна: {self.search_paths}")
        print(f"📂 Дэмжигдсэн файлын төрөл: {', '.join(SUPPORTED_EXTENSIONS)}")
        documents = []
        file_count = 0
        max_size_bytes = max_size_mb * 1024 * 1024
        for search_path in self.search_paths:
            if not os.path.exists(search_path):
                print(f"⚠️ Директори олдсонгүй: {search_path}")
                continue
            print(f"\n📁 Хайж байна: {search_path}")
            for root, dirs, files in os.walk(search_path):
                if self.should_skip_directory(root):
                    dirs[:] = []
                    continue
                for filename in files:
                    if file_count >= max_files:
                        print(f"\n⚠️ Хязгаарт хүрлээ: {max_files} файл")
                        break
                    filepath = os.path.join(root, filename)
                    ext = Path(filename).suffix.lower()
                    if ext not in SUPPORTED_EXTENSIONS:
                        continue
                    try:
                        file_size = os.path.getsize(filepath)
                        if file_size > max_size_bytes or file_size == 0:
                            continue
                    except:
                        continue
                    content = self.read_file(filepath)
                    if not content or len(content. strip()) < 50:
                        continue
                    file_hash = self.get_file_hash(filepath)
                    modified_time = datetime.fromtimestamp(os. path.getmtime(filepath))
                    doc = Document(
                        page_content=content[:5000],
                        metadata={
                            "filename": filename,
                            "filepath": filepath,
                            "extension": ext,
                            "size_kb": file_size / 1024,
                            "modified": modified_time.isoformat(),
                            "hash": file_hash
                        }
                    )
                    documents.append(doc)
                    file_count += 1
                    if file_count % 50 == 0:
                        print(f"✅ {file_count} файл боловсруулсан...")
                if file_count >= max_files:
                    break
        print(f"\n✅ Нийт {len(documents)} файл боловсруулсан")
        self.all_docs = documents
        return documents

    def create_index(self, documents):
        if not documents:
            print("❌ Индекс үүсгэх баримт байхгүй")
            return False
        print(f"\n🔄 FAISS индекс үүсгэж байна ({len(documents)} баримт)...")
        try:
            self.db = FAISS.from_documents(documents, self.embedding)
            os.makedirs(INDEX_FOLDER, exist_ok=True)
            self.db.save_local(INDEX_FOLDER)
            self.metadata = {
                "created": datetime.now().isoformat(),
                "num_documents": len(documents),
                "files": [doc.metadata for doc in documents]
            }
            with open(METADATA_FILE, 'wb') as f:
                pickle.dump(self.metadata, f)
            print(f"✅ Индекс амжилттай үүсгэгдлээ: {INDEX_FOLDER}")
            return True
        except Exception as e:
            print(f"❌ Индекс үүсгэхэд алдаа: {e}")
            return False

    def load_index(self):
        if not os.path.exists(INDEX_FOLDER):
            print("⚠️ Индекс олдсонгүй.  Эхлээд scan_disk() дуудна уу.")
            return False
        try:
            print("🔄 Индекс ачаалж байна...")
            self.db = FAISS.load_local(INDEX_FOLDER, self. embedding, allow_dangerous_deserialization=True)
            if os.path.exists(METADATA_FILE):
                with open(METADATA_FILE, 'rb') as f:
                    self. metadata = pickle.load(f)
                print(f"✅ Индекс ачаалагдлаа: {self.metadata. get('num_documents', 0)} баримт")
            else:
                print("⚠️ Metadata олдсонгүй")
            return True
        except Exception as e:
            print(f"❌ Индекс ачаалахад алдаа: {e}")
            return False

    def search_by_keyword(self, keyword):
        keyword = keyword.lower()
        results = []
        if self.metadata and "files" in self.metadata:
            for file_meta in self.metadata["files"]:
                if keyword in file_meta.get("filename", "").lower():
                    results.append(file_meta)
        return results

    def semantic_search(self, query, k=5, score_threshold=2.0):
        if not self.db:
            print("❌ Индекс ачаалаагүй байна")
            return []
        try:
            results = self.db.similarity_search_with_score(query, k=k)
            filtered = [(doc, score) for doc, score in results if score < score_threshold]
            if not filtered and results:
                print(f"⚠️ Threshold-оос давсан, бүх үр дүнг харуулж байна")
                filtered = results
            return filtered
        except Exception as e:
            print(f"❌ Хайлтад алдаа: {e}")
            return []

    def generate_answer(self, results, question, stream=True):
        """
        Ollama ашиглан хариулт үүсгэх
        """
        if not results:
            return "Холбогдох мэдээлэл олдсонгүй."
        
        if not self.llm.is_available():
            return "❌ Ollama ажиллахгүй байна.  'ollama serve' ажиллуулна уу."
        
        # Build context from search results
        snippets = []
        sources = []
        
        for i, (doc, score) in enumerate(results, 1):
            filename = doc.metadata.get("filename", f"source_{i}")
            filepath = doc.metadata.get("filepath", "")
            content = doc.page_content[:1500]. replace("\n", " ").strip()
            
            snippets.append(f"""
📄 Эх сурвалж {i}: {filename}
📂 Байршил: {filepath}
📝 Агуулга:
{content}
""")
            sources.append(filename)
        
        context = "\n---\n".join(snippets)
        
        # Create prompt
        prompt = f"""Та туслах AI байна.  Дараах эх сурвалжуудаас ЗӨВХӨН мэдээлэл ашиглан асуултад хариулна уу. 

📚 ЭХ СУРВАЛЖУУД:
{context}

❓ АСУУЛТ: {question}

📋 ЗААВАРЧИЛГАА:
1.  Зөвхөн өгөгдсөн эх сурвалжаас хариулна
2. Эх сурвалжийг [filename] хэлбэрээр дурдана
3. Хэрэв мэдээлэл байхгүй бол "Энэ мэдээлэл эх сурвалжид байхгүй байна" гэж хэлнэ
4. Товч, тодорхой хариулна
5. Асуулт ямар хэлээр байна тэр хэлээр хариулна

✍️ ХАРИУЛТ:"""

        if stream:
            # Streaming response
            print("\n💡 AI Хариулт:")
            full_response = ""
            for chunk in self.llm. generate_stream(prompt, temperature=0.3):
                print(chunk, end="", flush=True)
                full_response += chunk
            print()  # New line after streaming
            
            # Add sources if not mentioned
            if not any(s in full_response for s in sources):
                sources_text = f"\n\n📚 Ашигласан эх сурвалж: {', '.join(sources)}"
                print(sources_text)
                full_response += sources_text
            
            return full_response
        else:
            # Non-streaming response
            response = self.llm.generate(prompt, temperature=0.3, max_tokens=1024)
            
            if not any(s in response for s in sources):
                response += f"\n\n📚 Ашигласан эх сурвалж: {', '.join(sources)}"
            
            return response

    def interactive_search(self):
        print("\n" + "="*60)
        print("🧠 Диск Хайлтын AI Систем (Ollama)")
        print("="*60)
        print(f"🤖 Model: {self.llm.model}")
        print("💡 Командууд:")
        print("  - 'stats' - статистик харах")
        print("  - 'rescan' - дахин хайх")
        print("  - 'models' - моделууд харах")
        print("  - 'model <name>' - модел солих")
        print("  - 'exit' - гарах")
        print("="*60 + "\n")

        while True:
            try:
                user_input = input("🔍 Асуулт: ").strip()
            except (KeyboardInterrupt, EOFError):
                print("\n👋 Баяртай!")
                break
                
            if not user_input:
                continue
            if user_input.lower() == "exit":
                print("👋 Баяртай!")
                break
            if user_input.lower() == "stats":
                self.show_statistics()
                continue
            if user_input.lower() == "models":
                models = self.llm.list_models()
                print(f"\n📋 Боломжит моделууд: {', '.join(models)}")
                print(f"🔹 Одоогийн модел: {self.llm.model}\n")
                continue
            if user_input.lower(). startswith("model "):
                new_model = user_input[6:].strip()
                self.llm.model = new_model
                print(f"✅ Модел солигдлоо: {new_model}\n")
                continue
            if user_input.lower() == "rescan":
                print("\n🔄 Дахин хайж, индекс үүсгэж байна...")
                docs = self.scan_disk(max_files=500)
                self.create_index(docs)
                continue

            # Түлхүүр үгээр хайлт
            keyword_results = self.search_by_keyword(user_input)
            if keyword_results:
                print(f"\n[🔎 Түлхүүр үг] {len(keyword_results)} файл олдлоо:")
                for result in keyword_results[:3]:
                    print(f"  📄 {result['filename']} ({result['extension']})")
                    print(f"     📁 {result['filepath']}")

            # Утгын хайлт
            semantic_results = self.semantic_search(user_input, k=5)
            if not semantic_results:
                print("\n❌ Холбогдох мэдээлэл олдсонгүй.")
                print("💡 Зөвлөмж:")
                print("   - Өөр үг хэллэгээр оролдоно уу")
                print("   - 'stats' командаар ямар файлууд байгааг харна уу")
                print("   - 'rescan' командаар дахин хайж үзнэ үү\n")
                continue

            print(f"\n[📚 Утгын хайлт] {len(semantic_results)} баримт олдлоо:")
            for i, (doc, score) in enumerate(semantic_results, 1):
                snippet = doc.page_content[:200].replace("\n", " ") + "..."
                print(f"\n{i}. 🎯 Оноо: {score:.4f}")
                print(f"   📄 Файл: {doc.metadata. get('filename', 'Unknown')}")
                print(f"   📂 Зам: {doc.metadata.get('filepath', 'Unknown')}")
                print(f"   📏 Хэмжээ: {doc.metadata.get('size_kb', 0):.1f} KB")
                print(f"   📝 Агуулга: {snippet}")

            # Generate answer with Ollama (streaming)
            self.generate_answer(semantic_results[:3], user_input, stream=True)
            
            print("\n" + "-" * 60)

    def show_statistics(self):
        if not self.metadata:
            print("📊 Статистик байхгүй")
            return
        print("\n" + "="*60)
        print("📊 Системийн Статистик")
        print("="*60)
        print(f"📅 Үүссэн: {self.metadata. get('created', 'Unknown')}")
        print(f"📄 Нийт баримт: {self.metadata.get('num_documents', 0)}")
        print(f"🤖 Ollama Model: {self.llm.model}")
        print(f"🔌 Ollama Status: {'✅ Ажиллаж байна' if self.llm.is_available() else '❌ Ажиллахгүй'}")
        
        if "files" in self.metadata:
            extensions = {}
            total_size = 0
            for file_meta in self.metadata["files"]:
                ext = file_meta.get("extension", "unknown")
                size = file_meta.get("size_kb", 0)
                extensions[ext] = extensions.get(ext, 0) + 1
                total_size += size
            print(f"💾 Нийт хэмжээ: {total_size/1024:.2f} MB")
            print("\n📂 Файлын төрөл:")
            for ext, count in sorted(extensions.items(), key=lambda x: x[1], reverse=True):
                print(f"   {ext}: {count} файл")
            print("\n📋 Баримтын жагсаалт:")
            for i, file_meta in enumerate(self.metadata["files"][:20], 1):
                print(f"   {i}. {file_meta. get('filename', 'Unknown')} ({file_meta.get('extension', 'unknown')})")
            if len(self.metadata["files"]) > 20:
                print(f"   ... болон өөр {len(self.metadata['files']) - 20} файл")
        print("="*60 + "\n")


def main():
    print("🚀 Диск Хайлтын Систем (Ollama)\n")
    
    # Check Ollama
    print("🔍 Ollama шалгаж байна...")
    try:
        response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=5)
        if response.status_code == 200:
            models = [m["name"] for m in response.json(). get("models", [])]
            print(f"✅ Ollama бэлэн!  Моделууд: {', '.join(models[:5])}")
        else:
            print("⚠️ Ollama хариу өгөхгүй байна")
    except:
        print("❌ Ollama ажиллахгүй байна!")
        print("\n📋 Засах заавар:")
        print("   1. Шинэ PowerShell нээж: ollama serve")
        print("   2.  Model татсан эсэх: ollama list")
        print("   3. Model татаагүй бол: ollama pull qwen2")
        input("\nEnter дарж үргэлжлүүлэх...")
    
    print("\n📁 Хайх директоруудыг оруулна уу (таслалаар тусгаарлана):")
    print("   Жишээ: D:/Documents, D:/Projects, C:/Users/YourName/Desktop")
    print("   Хоосон орхивол D:/ болон C:/Users хайна")
    user_paths = input("\n📂 Директори: ").strip()
    if user_paths:
        search_paths = [p.strip() for p in user_paths.split(",")]
    else:
        search_paths = ["D:/", "C:/Users"]
    searcher = AdvancedDiskSearch(search_paths=search_paths)
    if os.path.exists(INDEX_FOLDER):
        print(f"\n✅ Хадгалсан индекс олдлоо: {INDEX_FOLDER}")
        choice = input("Ашиглах уу? (y/n): ").strip().lower()
        if choice == 'y':
            if searcher.load_index():
                searcher.interactive_search()
                return
    print("\n🔄 Диск хайж, индекс үүсгэж байна...")
    print("⚠️ Энэ удаан үргэлжлэх боломжтой (5-30 минут)")
    max_files = input("\nХамгийн их файлын тоо (default 1000): ").strip()
    max_files = int(max_files) if max_files. isdigit() else 1000
    docs = searcher.scan_disk(max_files=max_files, max_size_mb=10)
    if docs:
        if searcher.create_index(docs):
            searcher.interactive_search()
    else:
        print("❌ Файл олдсонгүй")


if __name__ == "__main__":
    main()